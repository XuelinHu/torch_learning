import os
from docx import Document

def replace_text_in_docx(file_path, replacements):
    doc = Document(file_path)
    for p in doc.paragraphs:
        for old in replacements:
            if old in p.text:
                p.text = p.text.replace(old, "")

    # 表格里的内容也要替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old in replacements:
                    if old in cell.text:
                        for p in cell.paragraphs:
                            p.text = p.text.replace(old, "")
    doc.save(file_path)

def batch_process_docx(folder):
    replacements = ["24-铁电45班", "23-铁电45班", "23-铁电46班"]
    for filename in os.listdir(folder):
        if filename.endswith(".docx") and not filename.startswith('.'):
            file_path = os.path.join(folder, filename)
            print(f"正在处理：{file_path}")
            replace_text_in_docx(file_path, replacements)
    print("处理完成 ✅")

if __name__ == "__main__":
    folder_path = '/Volumes/HIKSEMI/LT Work/202509教学工作/信息技术2-胡学林/教案'
    batch_process_docx(folder_path)
