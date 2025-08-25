from docx import Document
from sympy.strategies.core import switch

from tencent_client import batch_translate as tbt
from volcano_client import batch_translate as vbt
import re
from docx.shared import RGBColor
import json
import os
import shutil

import textract


def extract_all_tables(chapter,doc):
    tables = doc.tables
    text_list = []
    txt_map = {}
    idx_map = {
        1: '0_0_0',
        2: '0_2_1',
        3: '0_3_4',
        4: '1_1_3',
        5: '1_1_1',
        6: '1_1_2',
        7: '1_3_1',
        8: '1_3_1',
        9: '1_3_1',
        10: '1_4_1',
        11: '1_5_1',
    }
    idx = 1
    content_flag = False
    _13 = []
    _14 = []
    _15 = []

    for idx, table in enumerate(tables):
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 检查单元格是否是合并单元格的一部分
                paragraphs = [para.text.strip().replace('..', '') for para in cell.paragraphs if
                              para.text.strip()]

                for key, value in idx_map.items():
                    if "_".join([str(idx), str(i), str(j)]) == value:
                        if 1 == key:
                            txt_map[key] = chapter
                        if 2 == key:
                            txt_map[key] = paragraphs[0]
                        if 3==key:
                            txt_map[key]=paragraphs[0]
                        if 4==key:
                            txt_map[key]=paragraphs[0]
                        if 5==key:
                            txt_map[key]=paragraphs[0]
                        if 6==key:
                            txt_map[key]=paragraphs[0]
                        if 7==key:
                            txt_map[key]=paragraphs[2]
                        if 8==key:
                            txt_map[key]=paragraphs[4]
                        if 9==key:
                            txt_map[key]=paragraphs[6]
                        if 10==key:
                            txt_map[key]="\n".join(paragraphs)
                        if 11==key:
                            txt_map[key]="\n".join(paragraphs)

                if '课  后' in "".join(paragraphs) or '教学反思' in "".join(paragraphs):
                    txt_map[13] = "\n".join(_13[2:])
                    txt_map[14] = "\n".join(_14[2:])
                    txt_map[15] = "\n".join(_15[2:])
                    for k,v in txt_map.items():
                        print(k,v)
                    return txt_map
                if '课  中' in "".join(paragraphs):
                    content_flag = True
                if content_flag and j == 0:
                    for txt in paragraphs:
                        _13.append(txt)
                if content_flag and j == 2:
                    for txt in paragraphs:
                        _14.append(txt)
                if content_flag and j == 4:
                    for txt in paragraphs:
                        _15.append(txt)
                text_list.extend(paragraphs)

    return txt_map


def extract_paragraphs(doc):
    # 初始化段落列表
    paragraph_list = []

    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        # 提取段落文本（自动过滤空行）
        text = paragraph.text.strip()
        # 仅保留非空内容
        if text:
            paragraph_list.append(paragraph)
    idx = 0
    txt_set = set()
    while idx < len(paragraph_list):
        paragraph = paragraph_list[idx]
        text_strip = paragraph.text.strip()
        txt_set.add(text_strip)
        idx += 2
    else:
        idx += 1

    return paragraph_list, txt_set


def replace_text_in_docx(src_path, dst_path, replacements):
    """
    替换Word文档中的文本

    :param src_path: 源文档路径
    :param dst_path: 保存替换后的文档路径
    :param replacements: 替换规则字典，格式为 {待替换文本: 替换后文本}
    """
    # 检查源文件是否存在
    if not os.path.exists(src_path):
        raise FileNotFoundError(f"源文件不存在: {src_path}")

    # 打开文档
    doc = Document(src_path)

    # 1. 替换段落中的文本
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # 遍历段落中的所有运行元素（run）进行替换
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    # 2. 替换表格中的文本
    for table in doc.tables:
        # 遍历表格的每一行
        for row in table.rows:
            # 遍历每行的每个单元格
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, str(new_text))

    # 保存替换后的文档
    doc.save(dst_path)
    print(f"替换完成，已保存至: {dst_path}")

if __name__ == '__main__':

    def copy_with_template(folder_a, folder_b, template_file):
        # 确保目标文件夹存在
        os.makedirs(folder_b, exist_ok=True)

        # 遍历A中的文件
        idx = 1
        for filename in os.listdir(folder_a):
            if filename.lower().endswith(".docx") and not filename.lower().startswith("._"):  # 只处理 .doc 文件
                src_file = os.path.join(folder_a, filename)
                print(src_file)
                src_doc = Document(src_file)
                txt_map = extract_all_tables(idx,src_doc)
                idx+=1
                src = template_file
                dst = os.path.join(folder_b, filename)

                # 复制模板并重命名
                shutil.copy(src, dst)
                print(f"已生成: {dst}")
                replace_map = {}
                for k,v in txt_map.items():
                    _k = '${%s}' % k
                    replace_map[_k] = v
                print(replace_map)
                replace_text_in_docx(dst,dst,replace_map)

    if __name__ == "__main__":
        folder_b = r'/Volumes/HIKSEMI/LT Work/202509课件/电工基础-胡学林/教案-2024版'  # B 文件夹路径
        folder_a = r'/Volumes/HIKSEMI/LT Work/202509课件/电工基础-胡学林/教案-2024版-docx'
        template_file = r'/Volumes/HIKSEMI/LT Work/202509课件/电工基础-胡学林/教案-模板-代码版.doc'

        copy_with_template(folder_a, folder_b, template_file)
