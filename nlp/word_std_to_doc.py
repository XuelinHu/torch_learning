from docx import Document
from sympy.strategies.core import switch

from tencent_client import batch_translate as tbt
from volcano_client import batch_translate as vbt
import re
from docx.shared import RGBColor
import json
import os
import shutil

import textract



def replace_text_in_docx(src_path, dst_path, replacements):
    """
    替换Word文档中的文本

    :param src_path: 源文档路径
    :param dst_path: 保存替换后的文档路径
    :param replacements: 替换规则字典，格式为 {待替换文本: 替换后文本}
    """
    # 检查源文件是否存在
    if not os.path.exists(src_path):
        raise FileNotFoundError(f"源文件不存在: {src_path}")

    # 打开文档
    doc = Document(src_path)

    # 1. 替换段落中的文本
    for para in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in para.text:
                # 遍历段落中的所有运行元素（run）进行替换
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

    # 2. 替换表格中的文本
    for table in doc.tables:
        # 遍历表格的每一行
        for row in table.rows:
            # 遍历每行的每个单元格
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in para.text:
                            para.text = para.text.replace(old_text, str(new_text))

    # 保存替换后的文档
    doc.save(dst_path)
    print(f"替换完成，已保存至: {dst_path}")

    def copy_with_template(folder_source, folder_b, template_file):
        # 确保目标文件夹存在
        os.makedirs(folder_b, exist_ok=True)

        # 遍历A中的文件
        idx = 1
        for filename in sorted(os.listdir(folder_source)):
            if filename.lower().endswith(".docx") and not filename.lower().startswith("._"):  # 只处理 .doc 文件
                src_file = os.path.join(folder_source, filename)
                fname = filename.replace('docx', '')
                src_doc = Document(src_file)
                txt_map = extract_all_tables(fname, idx, src_doc)
                continue
                for k, v in txt_map.items():
                    print()
                    print(k, '===', v)

                idx += 1
                src = template_file
                dst = os.path.join(folder_b, filename)

                # 复制模板并重命名
                shutil.copy(src, dst)
                print(f"已生成: {dst}")
                replace_map = {}
                for k, v in txt_map.items():
                    _k = '${%s}' % k
                    replace_map[_k] = v
                print(replace_map)
                replace_text_in_docx(dst, dst, replace_map)


import re

def split_to_three(text: str):
    """
    将文本按；或。分段，组成一个长度为3的一维数组
    不够3个用空字符串补齐
    """
    # 按照中文分号、句号分割
    parts = re.split(r"[；。]", text)
    # 去掉空白
    parts = [p.strip() for p in parts if p.strip()]
    # 保留前3个，不足补空
    result = parts[:3] + [""] * (3 - len(parts))
    return result



def print_tables(docx_path, template_file):
    doc = Document(docx_path)
    idx_map = {
        2: 0,
        3: 1,
        4: 3,
        5: 2,
        6: 4,
        7: 5,
        8: 6,
    }

    prompt_command = """
此次课程是：${2}，请补充下面的内容，每一小项内容在80-120字，注意各小项内容的换行：
【上节课学情分析】
【教学重点的30字解决措施】:${6}
【教学重点的30解决措施】:${7}
【复习旧课】
【导入新课】 ‌案例展示 问题引导 导入主题 
【讲授新课】 
知识点1：${10}， 内容讲解 教学重点 教学难点
知识点1： 教师活动  学生活动  教学方法 
知识点2： ${11}，内容讲解  教学重点  教学难点 
知识点2： 教师活动 学生活动  教学方法 
知识点3： ${12}，内容讲解  教学重点 教学难点 
知识点3： 教师活动  学生活动  教学方法 
【课堂小结】 知识回顾 课程思政
    """

    idx = 1
    prompt_command_list = []
    for table_index, table in enumerate(doc.tables):  # 遍历表格
        if not 4 == table_index:
            continue
        for row_index, row in enumerate(table.rows):  # 遍历行
            if row_index < 2:
                continue
            pre_fname = ''
            replace_map = {}
            for col_index, cell in enumerate(row.cells):  # 遍历列
                text = cell.text.strip()
                idx_str = f"{table_index}_{row_index}_{col_index}"
                print(f"{idx_str}: {text}")
                for k, v in idx_map.items():
                    if col_index == v:
                        replace_map['${%d}' % k] = text
                if col_index ==0:
                    pre_fname = text
            arr = split_to_three(replace_map['${4}'])
            replace_map['${10}'] = arr[0]
            replace_map['${11}'] = arr[1]
            replace_map['${12}'] = arr[2]
            replace_map['${1}'] = str(idx)
            f_name = os.path.join(folder_target, pre_fname.replace("/","")+ '.docx')
            print(replace_map)
            print(f_name)
            # shutil.copy(template_file, f_name)
            # replace_text_in_docx(f_name,f_name,replace_map)
            idx+=1

            prompt_command_new = prompt_command
            for i in [2,6,7,10,11,12]:
                prompt_command_new = prompt_command_new.replace('${%d}' % i,replace_map['${%d}' % i])
            prompt_command_list.append(prompt_command_new)

    print("------------------")
    for i  in prompt_command_list:
        print(i)

if __name__ == '__main__':
    folder_source = '/Volumes/HIKSEMI/LT Work/202509教学工作/信息技术2-胡学林/《信息技术2》课程标准_v3.docx'
    folder_target = '/Volumes/HIKSEMI/LT Work/202509教学工作/信息技术2-胡学林/temp'
    template_file = '/Volumes/HIKSEMI/LT Work/202509教学工作/信息技术2-胡学林/教案/模版-代码版本.docx'

    print_tables(folder_source, template_file)
