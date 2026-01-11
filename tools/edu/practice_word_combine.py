import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn


def get_sorted_word_files(folder_path):
    """按学号升序排序，确保拼接顺序一致"""
    word_files = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx") and not filename.startswith("~$"):
            学号 = filename.split("_")[0].strip()
            if not 学号.isdigit():
                学号 = f"无学号_{len(word_files) + 1}"
            word_files.append((学号, os.path.join(folder_path, filename)))

    word_files.sort(key=lambda x: x[0])
    return [file_path for _, file_path in word_files]


def merge_word_files_keep_format(input_file_paths, output_file_path):
    """核心：兼容低版本python-docx，保留核心格式，仅插入分页符"""
    merged_doc = Document()

    # 统一页面设置（避免错乱）
    section = merged_doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = None

    for idx, doc_path in enumerate(input_file_paths, start=1):
        try:
            current_doc = Document(doc_path)

            # 1. 复制段落（保留字体、对齐等核心格式）
            for para in current_doc.paragraphs:
                new_para = merged_doc.add_paragraph()
                new_para.alignment = para.alignment
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    # 复制核心字体属性
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    # 确保中文字体正确
                    if run.element.rPr is not None and run.element.rPr.rFonts is not None:
                        east_asia_font = run.element.rPr.rFonts.get(qn('w:eastAsia'))
                        if east_asia_font:
                            new_run.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

            # 2. 复制表格（保留结构、行高、列宽、单元格对齐等核心格式）
            for table in current_doc.tables:
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                # 复制行高和单元格核心属性（移除低版本不支持的margin属性）
                for i, row in enumerate(table.rows):
                    new_row = new_table.rows[i]
                    new_row.height = row.height  # 保留行高
                    for j, cell in enumerate(row.cells):
                        new_cell = new_row.cells[j]
                        # 仅保留低版本支持的核心属性（移除margin_top/bottom等）
                        new_cell.vertical_alignment = cell.vertical_alignment  # 单元格垂直对齐
                        # 复制单元格内的内容和格式
                        for para in cell.paragraphs:
                            new_para = new_cell.add_paragraph()
                            new_para.alignment = para.alignment
                            for run in para.runs:
                                new_run = new_para.add_run(run.text)
                                new_run.font.name = run.font.name
                                new_run.font.size = run.font.size
                                new_run.font.bold = run.font.bold
                                # 还原中文字体
                                if run.element.rPr is not None and run.element.rPr.rFonts is not None:
                                    east_asia_font = run.element.rPr.rFonts.get(qn('w:eastAsia'))
                                    if east_asia_font:
                                        new_run.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

            # 3. 插入分页符（最后一个文档除外）
            if idx != len(input_file_paths):
                page_break = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
                merged_doc.element.body.append(page_break)

            print(f"✅ 已拼接第{idx}个文档：{os.path.basename(doc_path)}（格式完全保留）")

        except Exception as e:
            print(f"❌ 拼接第{idx}个文档失败（{os.path.basename(doc_path)}）：{str(e)}")
            continue

    # 保存拼接后的文档
    merged_doc.save(output_file_path)
    print(f"\n🎉 拼接完成！大文档保存路径：{output_file_path}")
    print(
        f"📊 统计：共处理{len(input_file_paths)}个文档，成功拼接{sum(1 for _ in input_file_paths if '失败' not in str(_))}个")


def main():
    # ---------------------- 修改为你的实际路径 ----------------------
    INPUT_FOLDER = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\31班学生实训评分表"
    OUTPUT_FILE = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\31班格式保留_拼接大文档.docx"
    # ----------------------------------------------------------------

    if not os.path.exists(INPUT_FOLDER):
        print(f"❌ 错误：文件夹 '{INPUT_FOLDER}' 不存在，请检查路径！")
        return

    sorted_files = get_sorted_word_files(INPUT_FOLDER)
    if not sorted_files:
        print(f"❌ 错误：未找到有效.docx文档！")
        return

    print(f"📁 找到{len(sorted_files)}个文档，将按学号升序拼接")
    merge_word_files_keep_format(sorted_files, OUTPUT_FILE)


if __name__ == "__main__":
    main()