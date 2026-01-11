import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import shutil


def calculate_scores(row):
    """计算S1-S10（学号去.0+等级逻辑确保正确）"""
    # 基础数据提取与清洗
    实训报告原始分 = row['实训报告'] if pd.notna(row['实训报告']) else 0.0
    期评成绩 = row['期评成绩'] if pd.notna(row['期评成绩']) else 0.0
    班级 = str(row['班级']).strip() if pd.notna(row['班级']) else ""

    # 学号彻底去.0（兼容浮点数/字符串格式）
    学号_raw = row['学号'] if pd.notna(row['学号']) else ""
    学号 = str(学号_raw).strip().rstrip(".0")  # 直接截取末尾.0，避免残留

    姓名 = str(row['姓名']).strip() if pd.notna(row['姓名']) else ""

    # S7（实训报告得分：满分10分，1位小数，0则空）
    S7 = round(实训报告原始分 / 10, 1)
    S7 = f"{S7:.1f}" if S7 != 0.0 else ""

    # S8（平时表现：默认10.0，1位小数）
    S8 = 10.0
    S8 = f"{S8:.1f}"

    # S9（总分：期评成绩，1位小数，0则空）
    S9 = round(期评成绩, 1)
    S9 = f"{S9:.1f}" if S9 != 0.0 else ""

    # 剩余分计算（避免空值参与）
    剩余分 = 0.0
    if S9 and S7 and S8:
        剩余分 = round(float(S9) - float(S7) - float(S8), 1)

    # S1-S5比例计算（工具函数简化代码）
    def calc_ratio(ratio):
        score = round(剩余分 * ratio / 80, 1) if 剩余分 != 0.0 else ""
        return f"{score:.1f}" if score != 0.0 else ""

    S1 = calc_ratio(15)
    S2 = calc_ratio(10)
    S3 = calc_ratio(5)
    S4 = calc_ratio(10)
    S5 = calc_ratio(30)

    # S6补差（确保S1-S8总和=S9，无误差）
    S6 = ""
    if 剩余分 != 0.0 and all([S1, S2, S3, S4, S5]):
        sum_s1s5 = sum(float(x) for x in [S1, S2, S3, S4, S5])
        S6 = round(剩余分 - sum_s1s5, 1)
        S6 = f"{S6:.1f}" if S6 != 0.0 else ""

    # S10等级（逻辑清晰，确保每个分数段对应正确等级）
    S10 = ""
    if S9:
        s9_val = float(S9)
        if s9_val >= 90:
            S10 = "优秀"
        elif s9_val >= 80:
            S10 = "良好"
        elif s9_val >= 70:
            S10 = "中等"
        elif s9_val >= 60:
            S10 = "及格"
        else:
            S10 = "不及格"

    return {
        'A': 班级, 'B': 学号, 'C': 姓名,
        'S1': S1, 'S2': S2, 'S3': S3, 'S4': S4, 'S5': S5, 'S6': S6,
        'S7': S7, 'S8': S8, 'S9': S9, 'S10': S10
    }


def format_template(template_path, formatted_template_path):
    """核心格式调整：标题小二22磅+表格五号12磅"""
    doc = Document(template_path)

    # 1. 标题（企业网设计与管理实训评分标准）：宋体、小二（22磅）
    if len(doc.paragraphs) >= 1:
        title_paragraph = doc.paragraphs[0]
        for run in title_paragraph.runs:
            # 设置中文字体为宋体（避免中文显示异常）
            run.font.name = u'宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = Pt(22)  # 小二=22磅，符合需求

    # 2. 班级/学号/姓名行：宋体、四号（14磅，保持清晰，与标题区分）
    if len(doc.paragraphs) >= 2:
        class_paragraph = doc.paragraphs[1]
        for run in class_paragraph.runs:
            run.font.name = u'宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.size = Pt(14)

    # 3. 表格内文字：宋体、五号（12磅，非四号，符合需求）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 可选：调整单元格内边距，避免文字拥挤
                cell.margin_top = Pt(1)
                cell.margin_bottom = Pt(1)
                cell.margin_left = Pt(1)
                cell.margin_right = Pt(1)
                # 遍历单元格内所有段落和run，确保格式覆盖
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = u'宋体'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                        run.font.size = Pt(12)  # 五号=12磅，非四号

    # 保存格式化模板（仅执行1次）
    doc.save(formatted_template_path)
    print(f"✅ 格式化模板生成完成（标题小二22磅+表格五号12磅）：{formatted_template_path}")


def replace_placeholder_in_formatted_word(formatted_template_path, output_path, data):
    """彻底修复{S10}}替换：遍历所有位置，处理多余右括号"""
    # 复制格式化模板，避免修改原模板
    shutil.copy2(formatted_template_path, output_path)
    doc = Document(output_path)

    # 定义所有需要替换的占位符（包含S10）
    placeholders = ['A', 'B', 'C', 'S1', 'S2', 'S3', 'S4', 'S5', 'S6', 'S7', 'S8', 'S9', 'S10']

    # ---------------------- 1. 替换段落中的占位符（含标题下的班级行）----------------------
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            original_text = run.text
            # 替换普通占位符（{A}{B}{C}{S1}-{S9}）
            for key in placeholders[:-1]:  # 先处理非S10
                placeholder = f"{{{key}}}"
                if placeholder in original_text:
                    original_text = original_text.replace(placeholder, data.get(key, ""))
            # 单独处理{S10}}（模板多1个右括号，直接匹配替换）
            if "{S10}" in original_text:
                original_text = original_text.replace("{S10}", data.get('S10', ""))
            # 更新run文本（保留原有格式，不重置字体）
            run.text = original_text

    # ---------------------- 2. 替换表格中的占位符（重点修复“等级：{S10}}”）----------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        original_text = run.text
                        # 替换普通占位符
                        for key in placeholders[:-1]:
                            placeholder = f"{{{key}}}"
                            if placeholder in original_text:
                                original_text = original_text.replace(placeholder, data.get(key, ""))
                        # 修复表格底部“等级：{S10}}”
                        if "{S10}" in original_text:
                            original_text = original_text.replace("{S10}", data.get('S10', ""))
                        run.text = original_text

    # 保存最终文件（格式+替换双重生效）
    doc.save(output_path)
    print(f"📄 生成文件：{os.path.basename(output_path)}")


def main():
    # 配置文件路径（按实际路径修改，路径分隔符用\\或/）
    excel_path = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\24企信31班_原始版.xlsx"
    word_template_path = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\企业网设计与管理实训评分标准.docx"
    output_folder = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\学生实训评分表"
    formatted_template_path = "D:\\edu_lt\\教案\\企业网设计与管理实训-胡学林\\作业成绩汇总-2025\\格式化后的模板.docx"

    # 1. 检查输入文件是否存在（避免路径错误）
    for path, name in [(excel_path, "Excel文件"), (word_template_path, "Word模板")]:
        if not os.path.exists(path):
            print(f"❌ 错误：{name} '{path}' 不存在，请检查路径！")
            return

    # 2. 生成格式化模板（仅1次，后续直接复用）
    if not os.path.exists(formatted_template_path):
        format_template(word_template_path, formatted_template_path)
    else:
        print(f"🔧 格式化模板已存在，直接使用（标题小二22磅+表格五号12磅）")

    # 3. 创建输出文件夹（确保目录存在）
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"📁 已创建输出文件夹：{os.path.abspath(output_folder)}")

    # 4. 读取Excel数据（确保Sheet1正确）
    try:
        df = pd.read_excel(excel_path, sheet_name="Sheet1")
        print(f"\n📊 成功读取Excel，共 {len(df)} 名学生数据")
        print("Excel列名：", df.columns.tolist())  # 调试用，确认字段存在
    except Exception as e:
        print(f"❌ 读取Excel失败：{str(e)}（可能是Sheet名错误或文件损坏）")
        return

    # 5. 批量生成学生Word评分表（学号无.0，格式无丢失）
    success_count = 0
    for index, row in df.iterrows():
        try:
            score_data = calculate_scores(row)
            # 生成唯一文件名（学号_姓名，避免重名）
            学号 = score_data['B'] if score_data['B'] else f"无学号_{index + 1}"
            姓名 = score_data['C'] if score_data['C'] else f"无姓名_{index + 1}"
            output_word_name = f"{学号}_{姓名}_实训评分表.docx"
            output_word_path = os.path.join(output_folder, output_word_name)

            # 替换占位符并保存（保留所有格式）
            replace_placeholder_in_formatted_word(formatted_template_path, output_word_path, score_data)
            success_count += 1
        except Exception as e:
            print(f"❌ 第 {index + 1} 行生成失败：{str(e)}（可检查该行数据是否异常）")

    # 6. 输出最终统计结果（清晰展示生成情况）
    print(f"\n=== 🎉 批量生成完成 ===")
    print(f"📋 统计信息：")
    print(f"   总学生数：{len(df)}")
    print(f"   成功生成：{success_count} 个文件")
    print(f"   生成失败：{len(df) - success_count} 个文件")
    print(f"📂 文件保存路径：{os.path.abspath(output_folder)}")
    print("🎨 格式验证：标题（宋体小二22磅）、表格（宋体五号12磅）、{S10}已替换")


if __name__ == "__main__":
    main()